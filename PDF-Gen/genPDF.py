import docx
import numpy as np
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.enum.table import WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from math import floor
from datetime import date
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import docx2pdf




days_map = {0: 'الإثنين', 1: "الثلاثاء", 2: "الأربعاء", 3: "الخميس", 4: "الجمعة", 5: "السبت", 6:"الأحد"}

arabic_numbers={'0':'٠', '1':'١', '2': '٢', '3':'٣', '4':'٤', '5':'٥', '6':'٦', '7':'٧', '8':'٨', '9':'٩'}




def translate_all_numbers_to_arabic(English_text):
    arabicText=list(English_text)
    for i,c in enumerate(English_text):
        if c in list('0123456789'):
            arabicText[i] = arabic_numbers[c]
        elif c == '-':
            arabicText[i] = '/'
        else:
            arabicText[i] = c
    
    return ''.join(arabicText)

def cleanse_styles(doc):
    for style in doc.styles:
        print(style.name)
        if style.builtin:
            style.delete()

def Generate_A4_With_Header_And_footer(logo_path, background_image, centered_text, department_name, date_of_document, footer_person, footer_job):
    testDoc = docx.Document('../data/watermarked.docx')
    # new_section = testDoc.add_section(WD_SECTION.CONTINUOUS)
    section = testDoc.sections[0]

    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.3)
    section.page_height = Inches(11.7)

    section.left_margin = Inches(0.1)
    section.right_margin = Inches(0.1)
    section.top_margin = Inches(0.05)
    section.bot_margin = Inches(0.1)


    # header = section.header
    # paragraph = header.paragraphs[0]
    # run = paragraph.add_run()
    #testDoc.add_picture(background_image, width=Inches(8.3), height=Inches(11.7))


    docHeader = section.header
    table = docHeader.add_table(rows=1, cols=3, width=Inches(8.3))
    table.allignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.height=Pt(2)
    # table.autofit=True
    table.style = 'Header Table'

    

    Logo_Cell = table.cell(0, 0)
    logo_paragraph = Logo_Cell.add_paragraph()
    #logo_paragraph.style = testDoc.styles["Heading 1 Char"]
    newRun = logo_paragraph.add_run()
    newRun.add_picture (logo_path, width=Inches(2.5))

    #cleanse_styles(testDoc)
    Centered_table_style = testDoc.styles.add_style('Table big', WD_STYLE_TYPE.PARAGRAPH)
    Centered_table_style.base_style = testDoc.styles['Normal']

    centered_text_cell = table.cell(0, 1)
    centered_text_paragraph = centered_text_cell.add_paragraph(text= centered_text, style=Centered_table_style)
    centered_text_paragraph.style.font.size = Pt(15)
    centered_text_paragraph.style.font.bold = True
    centered_text_paragraph.style.font.color.rgb = RGBColor(0x0, 0x0, 0x00)
    
    paragraph_format = centered_text_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    table_style = testDoc.styles.add_style('Table Normal', WD_STYLE_TYPE.PARAGRAPH)
    table_style.base_style = testDoc.styles['Normal']

    right_hand_text_cell_0 = table.cell(0, 2)
    right_hand_text_cell_0.text = f"""جمهورية مصر العربية
    وزارة الدفاع
    {department_name}
    بتاريخ: {date_of_document}"""
    right_hand_text_cell_0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # right_hand_text_cell_2 = table.cell(2, 2)
    # right_hand_text_cell_2.text= "جهاز مستقبل مصر للتنمية المستدامة"
    # right_hand_text_cell_2.style=table_style
    # right_hand_text_cell_2.style.font.size = Pt(11)
    # right_hand_text_cell_2.style.font.bold = True
    # right_hand_text_cell_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # under_style = testDoc.styles.add_style('Table Normal Underlined', WD_STYLE_TYPE.PARAGRAPH)
    # under_style.base_style = testDoc.styles['Normal']

    # right_hand_text_cell_3 = table.cell(3, 2)
    # right_hand_text_cell_3.text= department_name
    # right_hand_text_cell_3.style=under_style
    # right_hand_text_cell_3.style.font.size = Pt(11)
    # right_hand_text_cell_3.style.font.bold = True
    # right_hand_text_cell_3.style.font.underline = True
    # right_hand_text_cell_3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # right_hand_text_cell_4 = table.cell(4, 2)
    # right_hand_text_cell_4.text= f"التاريخ: {date_of_document}"
    # right_hand_text_cell_4.style=under_style
    # right_hand_text_cell_4.style.font.size = Pt(11)
    # right_hand_text_cell_4.style.font.bold = True
    # right_hand_text_cell_4.style.font.underline = True
    # right_hand_text_cell_4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER



    docFooter = section.footer
    par = docFooter.add_paragraph(text=f'''{footer_person}\n{footer_job}''')
    par.style = testDoc.styles['Normal']
    par.style.font.color.rgb= RGBColor(0,0,0)
    par.style.font.bold=True
    par.style.font.size=Pt(20)
    


    return testDoc
# headerParagraph.




def addTitle(document, text, color):
    center_paragraph = document.add_paragraph(text)
    paragraph_format = center_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    [print(c.name) for c in document.styles]
    center_paragraph.style = document.styles['Heading 1']
    center_paragraph.style.font.size = Pt(26)
    center_paragraph.style.font.underline = True
    center_paragraph.style.font.color.rgb = color





def addArabicText(document, text, underline:bool, size, color_RGB):
    new_par = document.add_paragraph(text)
    paragraph_format = new_par.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    new_par.style = document.styles['Heading 1']
    new_par.style.font.size = size
    new_par.style.font.underline = underline
    new_par.style.font.color.rgb = RGBColor(*color_RGB)








def insertCharInsideString(sourceString, addedChar, index):
    return sourceString[:index+1] + addedChar + sourceString[index+1:]

def addDummyChars(text:str, text_size, bounding_inches):
    unlinkable_chars = list('ةاذدور زأإءآىؤى')
    middle_index = floor(len(text) / 2)
    distance_from_middle = 0
    while(text[middle_index + distance_from_middle] in unlinkable_chars and text[middle_index - distance_from_middle] in unlinkable_chars):
        distance_from_middle += 1
    

    non_preceded_chars = list(' ء')
    additional_distance = 1
    print(text)
    while(text[middle_index + distance_from_middle + additional_distance] in non_preceded_chars and text[middle_index - distance_from_middle - additional_distance] in non_preceded_chars):
            print(text[middle_index + distance_from_middle + additional_distance])
            additional_distance += 1
            if(text[middle_index + distance_from_middle + additional_distance] == 'ل' and text[middle_index + distance_from_middle + additional_distance+1] in list("أ ا إ آ")):
                 additional_distance += 1
    
    additional_distance = 0 if additional_distance == 1 else additional_distance

    targetIndex = (middle_index + distance_from_middle + additional_distance) if text[middle_index + distance_from_middle + additional_distance] not in unlinkable_chars else (middle_index - distance_from_middle - additional_distance)


    while(Pt(text_size)*len(text) < Inches(bounding_inches+5)):
        text = insertCharInsideString(text, 'ـ', targetIndex)

    return text










def set_table_header_bg_color(cell, color):
    """
    set background shading for Header Rows
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    constructed_Hex_color = ''
    for c in color:
        hex_str = str(hex(c))[2:]
        if(len(hex_str) == 1):
            hex_str = '0' + hex_str
        constructed_Hex_color = constructed_Hex_color + hex_str

    xml_string = r'<w:shd {} w:fill="'+constructed_Hex_color+'"/>'
    
    shading_elm = parse_xml(xml_string.format(nsdecls('w')))
    tblCellProperties.append(shading_elm)
    return cell




def addTableFormatted(document, num_rows, num_cols, headers, input_rows, fill_entire_width:bool, header_color_rgb= None, header_font_size=24, cells_font_size = 20, font_color = (0, 0, 0)):
    #some assertions
    for c in header_color_rgb:
        if c > 255 or c < 0:
            raise TypeError
    for c in font_color:
        if c > 255 or c < 0:
            raise TypeError
    
    
    table = document.add_table(rows=num_rows, cols=num_cols)
    # table.autofit=True
    table.style = 'Table Grid'
    table.table_direction=WD_TABLE_DIRECTION.RTL
    table.allignment = WD_TABLE_ALIGNMENT.RIGHT


    #decorating the header row and filling it
    for row in table.rows:
        row.height = Pt(header_font_size)

    while (len(headers) > num_cols):
        table.add_column()

    for i, header in enumerate(headers):
        par = table.cell(0, i).add_paragraph(header)
        paragraph_format = par.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.style.font.size = Pt(header_font_size)
        par.style.font.bold = True
        par.style.font.underline = False
        table.cell(0,i)._element = set_table_header_bg_color(table.cell(0,i), header_color_rgb)
        par.style.font.color.rgb = RGBColor(*font_color) if font_color else RGBColor(0,0,0)

        #resizing the cell to combat large text in smaller cells
        if(par.style.font.size.inches * len(header) > table.cell(0, i).width):
            table.cell(0, i).width = Inches(par.style.font.size.inches * len(header))

        #modifying the text to combat small text in larger cells, by adding dummy "ـ"
        par.text = addDummyChars(header, par.style.font.size.pt, table.cell(0, i).width.inches)

    for i, input_row in enumerate(input_rows):
        for j, newcell in enumerate(table.row_cells(i+1)): #i+1 so as to skip the first row, aka the header row
                # newcell.text = input_rows[i][j]
                par = newcell.add_paragraph(input_rows[i][j])
                newcell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                par.style.font.size = Pt(cells_font_size)





def ConvertAndSave(document, name):
    document.save(f'{name}.docx')
    docx2pdf.convert(f'{name}.docx', f"{name}.pdf")





docc = Generate_A4_With_Header_And_footer(logo_path='../data/logo.png', background_image='../data/logo_bg.png',centered_text="بسم الله الرحمن الرحيم", department_name="إدارة المتابعة", date_of_document=date.today().isoformat(), footer_person="المقدم / أحمد جمال خليفة", footer_job="مدير إدارة المتابعة")

addTitle(docc, text=f"بيان بالتمام عن اليوم {days_map[date.today().weekday()]} الموافق {translate_all_numbers_to_arabic(date.today().isoformat())}", color=RGBColor(0,0,0))

addArabicText(document=docc, text="أولاً: ضباط الصف", underline=True, size=Pt(13), color_RGB=(0,0,0))
addTableFormatted(document=docc, num_rows=2, num_cols=3, headers=['القوة', 'موجود', 'الغياب'], input_rows=(['1', '1', '0'],), fill_entire_width=True, header_color_rgb=(0xd9, 0xd9, 0xd9), header_font_size=15, cells_font_size=13, font_color=(0,0,0))

addArabicText(document=docc, text="ثانياً: الجنود", underline=True, size=Pt(13), color_RGB=(0,0,0))
addTableFormatted(document=docc, num_rows=2, num_cols=3, headers=['القوة', 'موجود', 'الغياب'], input_rows=(['15', '7', '8'],), fill_entire_width=True, header_color_rgb=(0xd9, 0xd9, 0xd9), header_font_size=15, cells_font_size=13, font_color=(0,0,0))

ConvertAndSave(document=docc, name='output_Generated2')

del docc