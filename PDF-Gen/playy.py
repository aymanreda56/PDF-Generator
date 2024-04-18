import docx
import docx.document
import numpy as np
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.enum.table import WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from math import floor
from datetime import date
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import docx2pdf




days_map = {0: 'الإثنين', 1: "الثلاثاء", 2: "الأربعاء", 3: "الخميس", 4: "الجمعة", 5: "السبت", 6:"الأحد"}

arabic_numbers={'0':'٠', '1':'١', '2': '٢', '3':'٣', '4':'٤', '5':'٥', '6':'٦', '7':'٧', '8':'٨', '9':'٩'}


def translate_all_numbers_to_arabic(English_text):
    arabicText=list(English_text)
    for i,c in enumerate(English_text):
        if c in list('0123456789'):
            arabicText[i] = arabic_numbers[c]
        elif c == '-':
            arabicText[i] = '/'
        else:
            arabicText[i] = c
    
    return ''.join(arabicText)


Department_Name = "إدارة المتابعة"
Date_Arabic =translate_all_numbers_to_arabic(date.today().isoformat())
Weekday = days_map[date.today().weekday()]

number_officers = 1
present_officers = 1
absent_officers = 0

number_caps = 1
present_caps = 0
absent_caps = 1

number_soldiers = 15
present_soldiers = 7
absent_soldiers = 8


fields_to_replace = {"$DEPT_NAME$": Department_Name,
                     "$DATE_AR$": Date_Arabic,
                     "$WEEKDAY$": Weekday,
                     "$NUM_OFFIC$": number_officers,
                     "$PRES_OFFIC$":present_officers,
                     "$ABS_OFFIC$":absent_officers,
                     "$NUM_CAPS$":number_caps,
                     "$PRES_CAPS$":present_caps,
                     "$ABS_CAPS$":absent_caps,
                     "$NUM_SOLD$":number_soldiers,
                     "$PRES_SOLD$":present_soldiers,
                     "$ABS_SOLD$":absent_soldiers,
                     }

def replace_placeHolders(fields_to_replace, carrier_object):
    for k,v in fields_to_replace.items():
        carrier_object = carrier_object.replace(k, translate_all_numbers_to_arabic(str(v)))
    return carrier_object


doccc = docx.Document('../pdf/templates/tamam.docx')

def Replace_Placeholders_Inside_Document(doc, fields_to_replace:dict):
    # Iterate over paragraphs
   
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                
                

        for p in section.footer.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)


     
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
            

    # Iterate over tables
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                        
    

# Open the Word document
# Iterate over all text in the document
Replace_Placeholders_Inside_Document(doccc)

# def find_replace_text(doc, find_text, replace_text):
#     # Iterate over paragraphs
#     for paragraph in doc.paragraphs:
#         if find_text in paragraph.text:
#             for run in paragraph.runs:
#                 run.text = run.text.replace(find_text, replace_text)

#     # Iterate over tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if find_text in cell.text:
#                     for paragraph in cell.paragraphs:
#                         for run in paragraph.runs:
#                             run.text = run.text.replace(find_text, replace_text)


# find_replace_text(doccc, "$cell$", "Ayman")

# # Save the modified document
doccc.save("modified_document.docx")