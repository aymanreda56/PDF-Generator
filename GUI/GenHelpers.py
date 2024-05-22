import docx
import docx.document
# from docx.enum.section import WD_SECTION, WD_ORIENT
# from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.enum.table import WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
# from math import floor
from datetime import date
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml
import helpers
import enums
import docx2pdf
import PyPDF2
import os, re




days_map = {0: 'الإثنين', 1: "الثلاثاء", 2: "الأربعاء", 3: "الخميس", 4: "الجمعة", 5: "السبت", 6:"الأحد"}

arabic_numbers={'0':'٠', '1':'١', '2': '٢', '3':'٣', '4':'٤', '5':'٥', '6':'٦', '7':'٧', '8':'٨', '9':'٩'}


def translate_all_numbers_to_arabic(English_text):
    arabicText=list(English_text)
    for i,c in enumerate(English_text):
        if c in list('0123456789'):
            arabicText[i] = arabic_numbers[c]
        elif c == '-':
            arabicText[i] = '/'
        else:
            arabicText[i] = c

    #fixing the date direction:
    arabicText = ''.join(arabicText)
    # print(arabicText)
    # print(English_text)
    for i in English_text:
        if i not in list('0123456789-/'): #test if there are any characters, ie; not a date format
            return arabicText

    # then fix the date format
    try:
        arabicText = re.split('/', arabicText)
        if(len(arabicText[0]) == 2 and len(arabicText[2]) == 4):
            arabicText.reverse()
        arabicText = '/'.join(arabicText)
    except:
        arabicText = ''.join(arabicText)
    # print(arabicText)
    # print('\n\n\n\n\n\n')
    return arabicText


Department_Name = "مكتب السيد/ مدير الجهاز"
Vacation_Begin_Hour = '900'
Vacation_End_Hour = '1000'
Date_Arabic =translate_all_numbers_to_arabic(date.today().isoformat())
Weekday = days_map[date.today().weekday()]




def change_style_of_par(doc, parapgraph, style_name: str, size: int, bold: bool):
    parapgraph.style = doc.styles[style_name]
    parapgraph.style.font.size = Pt(size)
    parapgraph.style.font.ltr = True
    parapgraph.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    parapgraph.style.font.bold = bold





def replace_placeHolders(fields_to_replace, carrier_object):
    for k,v in fields_to_replace.items():
        carrier_object = carrier_object.replace(k, translate_all_numbers_to_arabic(str(v)))
    return carrier_object


def Replace_Placeholders_Inside_Document(doc, fields_to_replace:dict, Iterable_Fields):
    # Iterate over paragraphs
   
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                
                

        for p in section.footer.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)


    space_paragraph = None
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
            if('$space$' in run.text):
                space_paragraph = paragraph
    
            
            

    # Iterate over tables
    Vacations_Table = False
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                        # print(f'\n\n\n\n{run.text}')
                        if("$SIGN$" in run.text):
                            run.text = re.sub('\(*\)*\s*\$SIGN\$\s*\)*\(*', '', run.text)
                            run.add_picture('../data/signature.png', width=Inches(1))
                            cell.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.keep_together = True 
                        if('$LOGO$' in run.text):
                            run.text = ''
                            run.add_picture('../data/logo.png', width=Inches(1.5))
                            cell.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if('$NUM$' in run.text or Vacations_Table == table):
                            Vacations_Table = table
                            break
    
    if(Vacations_Table and len(Iterable_Fields)):
        num_Vacations = len(Iterable_Fields)
        for i in range(num_Vacations - 1):
            Vacations_Table.add_row()
        
        counter = 1
        page_broken = False
        
        for i, row in enumerate(Vacations_Table.rows):
            
            if(i < 2):
                continue

            j= i-2
            
            row.cells[5].paragraphs[0].text = translate_all_numbers_to_arabic(Iterable_Fields[j]['To_Date'])
            change_style_of_par(doc, row.cells[5].paragraphs[0], 'Strong Par Small', 13, False)
            row.cells[4].paragraphs[0].text = translate_all_numbers_to_arabic(Iterable_Fields[j]['From_Date'])
            change_style_of_par(doc, row.cells[4].paragraphs[0], 'Strong Par Small', 13, False)

            row.cells[3].paragraphs[0].text = "أجازة ميدانية"
            change_style_of_par(doc, row.cells[3].paragraphs[0], 'Strong Par Small', 14, True)

            row.cells[1].paragraphs[0].text = Iterable_Fields[j]['Level']
            change_style_of_par(doc, row.cells[1].paragraphs[0], 'Strong Par', 14, True)

            row.cells[2].paragraphs[0].text = Iterable_Fields[j]['Name']
            change_style_of_par(doc, row.cells[2].paragraphs[0], 'Strong Par', 14, True)


            row.cells[0].paragraphs[0].text = translate_all_numbers_to_arabic(str(counter))
            change_style_of_par(doc, row.cells[0].paragraphs[0], 'Strong Par', 14, True)

            counter += 1
            if(i > 9 and (not page_broken)):
                par_format = row.cells[0].paragraphs[0].paragraph_format
                par_format.page_break_before = True
                page_broken = True
        
        
        space_paragraph.text = ''
        if counter < 8:
            space_paragraph.style.font.size = Pt((8-counter) * 9)
        

        # for row in Vacations_Table.rows:
        #     for cell in row.cells:
        #         cell.paragraphs[0].style = doc.styles['Strong Par']
        #         cell.paragraphs[0].style.font.size = Pt(16)
        #         cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #         cell.paragraphs[0].style.font.bold = True
    else:
        print('NO MORE VACATIONS')
        new_merged_cell = Vacations_Table.cell(2,0).merge(Vacations_Table.cell(2,5))
        for p in new_merged_cell.paragraphs:
            new_merged_cell._element.remove(p._element)
        new_merged_cell.add_paragraph(text = 'لا يــوجـد خــوارج')
        change_style_of_par(doc, new_merged_cell.paragraphs[0], 'Strong Par', 18, True)
        # allnewtables = []
        # for table in doc.tables:
        #     if table == Vacations_Table:
        #         doc.remove(table)







def Replace_Placeholders_Inside_Movements(doc, fields_to_replace:dict, Iterable_Fields):
    # Iterate over paragraphs
   
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                
                

        for p in section.footer.paragraphs:
            for run in p.runs:
                run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)



    space_paragraph = None
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
            if('$space$' in run.text):
                space_paragraph = paragraph

            

    # Iterate over tables
    Vacations_Table = False
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = replace_placeHolders(fields_to_replace=fields_to_replace, carrier_object=run.text)
                        if('$LOGO$' in run.text):
                            run.text = ''
                            run.add_picture('../data/logo.png', width=Inches(2.5))
                            cell.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if('$NUM$' in run.text or Vacations_Table == table):
                            Vacations_Table = table

    
    if(Vacations_Table and len(Iterable_Fields)):
        num_Vacations = len(Iterable_Fields)
        for i in range(num_Vacations-1):
            Vacations_Table.add_row()
        
        counter = 1
        for i, row in enumerate(Vacations_Table.rows):
            
            if(i < 2):
                continue

            j= i-2
            row.cells[4].paragraphs[0].text = translate_all_numbers_to_arabic(Iterable_Fields[j]['$TODATE$'])
            row.cells[3].paragraphs[0].text = translate_all_numbers_to_arabic(Iterable_Fields[j]['$FROMDATE$'])
            row.cells[1].paragraphs[0].text = Iterable_Fields[j]['$LEVEL$']
            row.cells[2].paragraphs[0].text = Iterable_Fields[j]['$NAME$']
            row.cells[0].paragraphs[0].text = translate_all_numbers_to_arabic(str(counter))
            counter += 1

        for row in Vacations_Table.rows:
            for cell in row.cells:
                cell.paragraphs[0].style = doc.styles['Strong Par']
                cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].style.font.bold = True
    
    
        space_paragraph.text = '  '
        sample_text = ''
        if counter < 16:
            loop_stop_counter = 16-counter
            for i in range (loop_stop_counter):
                sample_text = sample_text + '\n'
        space_paragraph.text = space_paragraph.text + sample_text
        
    
    else:
        allnewtables = []
        for table in doc.tables:
            if table == Vacations_Table:
                doc.remove(table)

        





def Replace_Placeholders_Inside_Vac_Passes(doc, fields_to_replace:list, Iterable_Fields:list):
    # Iterate over tables    
    counter = 0
    truncate_rest = False
    for table in doc.tables:
        final_row_idx = 0
        for row_index, row in enumerate(table.rows):
            if final_row_idx == 0 and len(Iterable_Fields):
                for i, tb_cell in enumerate(row.cells):
                    if(counter < len(Iterable_Fields)):
                        for tb in tb_cell.tables:
                            for tb_row in tb.rows:
                                for cell in tb_row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            for field in fields_to_replace:
                                                    run.text = run.text.replace(field, Iterable_Fields[counter][field])
                                                    # print(Iterable_Fields[counter][field])
                                                    


                        for paragraph in tb_cell.paragraphs:
                                        for run in paragraph.runs:
                                            for field in fields_to_replace:
                
                                                    run.text = run.text.replace(field, Iterable_Fields[counter][field])
                                                    # print(Iterable_Fields[counter][field])
                                                    
                                                
                        counter += 1
                    else:
                        for tbb in tb_cell.tables:
                            for r in tbb.rows:
                                for tbcell in r.cells:
                                    for paragraph in tbcell.paragraphs:
                                        for run in paragraph.runs:
                                            run.text = ''
                        for paragraph in tb_cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = ''
                        if final_row_idx == 0 and len(Iterable_Fields):
                            final_row_idx = row_index
            else:
                break
                        
            
        for idx in range(len(table.rows) -1 , final_row_idx, -1):
                table._element.remove(table.rows[idx]._tr)
        

                    # tb_cell._element.getparent().remove(tb_cell._element)
                    # for idx in range(len(table.rows) -1 , row_index, -1):
                    #     table._element.remove(table.rows[idx]._tr)
                    # row._element.remove(tb_cell._element)
                    
                    # for tbb in tb_cell.tables:
                    #     for r in tbb.rows:
                    #         for tbcell in r.cells:
                    #             # tbb._element.remove(tbcell._element)
                    #             for paragraph in tbcell.paragraphs:
                    #                 for run in paragraph.runs:
                    #                     run.text = ''
                    # for paragraph in tb_cell.paragraphs:
                    #                 for run in paragraph.runs:
                    #                     run.text = ''
                                        
                



def Export_Movements_Aggrigator():
    AllVacations = helpers.getActiveVacations(with_disabled=False, only_extensions=0)
    ExtendedVacations = helpers.getActiveVacations(with_disabled=False, only_extensions=1)
    dictionary_For_Different_Dates = {}
    allDates = helpers.getAllDates(only_extensions=0)
    for i in allDates:
        dictionary_For_Different_Dates[(i,0)] = []

    ExtendedDates = helpers.GetAllFromDatesOfExtensions()
    for i in ExtendedDates:
        dictionary_For_Different_Dates[(i,1)] = []

    for i in AllVacations:
        dictionary_For_Different_Dates[(i[2], 0)].append(i)

    for i in ExtendedVacations:
        modified_entry = list(i)
        modified_entry[2] = helpers.GetExtensionFromDate(i[0]) 
        dictionary_For_Different_Dates[(helpers.GetExtensionFromDate(i[0]), 1)].append(modified_entry)
    
    keys_to_be_deleted = []
    for k,v in dictionary_For_Different_Dates.items():
        if(not v):
            keys_to_be_deleted.append(k)
    
    for k in keys_to_be_deleted:
        del dictionary_For_Different_Dates[k] 

    return dictionary_For_Different_Dates



def Export_Movements_PDF_For_One_Date(arr_of_entries:list, date_of_vac:str, extended:int = 0):
    Date_Arabic =translate_all_numbers_to_arabic(date_of_vac)
    today_Date_Arabic = translate_all_numbers_to_arabic(date.today().isoformat())
    Weekday = days_map[date.fromisoformat(date_of_vac).weekday()]
    # AllVacations = helpers.getActiveVacations(with_disabled=False)
    IterableFieldsDict = []
    for i, tup in enumerate(arr_of_entries):
        sampleDict = {}
        sampleDict['Soldier_ID'] = tup[0]
        sampleDict['$NAME$'] = tup[1]
        sampleDict['$FROMDATE$'] = tup[2]
        sampleDict['$TODATE$'] = tup[3]
        sampleDict['$STATE$'] = tup[4]
        sampleDict['$SUMMONED$'] = tup[5]
        sampleDict['$LEVEL$'] = enums.ArmyLevels[int(helpers.getLevelFromID(tup[0]))-1]
        sampleDict["$NUM$"] = i
        sampleDict["$DEPT_NAME$"]= Department_Name
        sampleDict["$WEEKDAY$"] = Weekday
        sampleDict["$DATE_AR$"]=Date_Arabic
        IterableFieldsDict.append(sampleDict)

    fields_to_replace = {"$DEPT_NAME$":Department_Name, '$DATE_AR$':Date_Arabic, '$WEEKDAY$':Weekday, "$TOD_AR$":today_Date_Arabic, '$DOC_TYP$': 'إمتداد' if extended else 'تحركات'}
    doccc = docx.Document('../pdf/templates/movements.docx')
    Replace_Placeholders_Inside_Movements(doccc, fields_to_replace, Iterable_Fields=IterableFieldsDict)
    ConvertAndSave(document=doccc, typeDoc='يوميات_تحركات', date_of_doc=date_of_vac, put_watermark=True)


def Export_Movements_PDF():
    all_entries_dict = Export_Movements_Aggrigator()
    for k,v in all_entries_dict.items():
        Export_Movements_PDF_For_One_Date(v, k[0], k[1])



def Export_Vacation_Passes_PDF():
    AllVacations = helpers.getActiveVacations(with_disabled=False)

    IterableFieldsDict = []
    for i, tup in enumerate(AllVacations):
        sampleDict = {}
        sampleDict['$SOLDNAME$'] = tup[1]
        sampleDict["$FROM_DATE$"] = translate_all_numbers_to_arabic(tup[2])
        sampleDict["$TO_DATE$"] = translate_all_numbers_to_arabic(tup[3])
        sampleDict['$LEVEL$'] = "جندي" if (enums.ArmyLevels[int(helpers.getLevelFromID(tup[0]))-1] == 'عسكري') else enums.ArmyLevels[int(helpers.getLevelFromID(tup[0]))-1]
        sampleDict["$BH$"] = translate_all_numbers_to_arabic(Vacation_Begin_Hour)
        sampleDict["$EH$"] = translate_all_numbers_to_arabic(Vacation_End_Hour)
        sampleDict['$NUM$'] = translate_all_numbers_to_arabic(str(i+1))
        IterableFieldsDict.append(sampleDict)

    fields_to_replace = ["$SOLDNAME$",
                        "$FROM_DATE$",
                        "$TO_DATE$", "$LEVEL$", "$BH$", "$EH$"]

    doccc = docx.Document('../pdf/templates/vacation_passes.docx')
    Replace_Placeholders_Inside_Vac_Passes(doccc, fields_to_replace=fields_to_replace, Iterable_Fields = IterableFieldsDict)
    ConvertAndSave(document=doccc, typeDoc='تصاريح', put_watermark=False)




def Export_Tamam_PDF():

    number_officers = int(helpers.getNumOfficers())
    absent_officers = len(helpers.getAbsentOfficers())
    present_officers = number_officers - absent_officers



    number_caps = int(helpers.getNumCaps())
    absent_caps = len(helpers.getAbsentCaps())
    present_caps = number_caps - absent_caps


    number_soldiers = int(helpers.getNumSoldiers())
    absent_soldiers = len(helpers.getAbsentSoldiers())
    # print(helpers.getAbsentSoldiers())
    present_soldiers = number_soldiers - absent_soldiers
    # print(f'Num Soldiers = {number_soldiers}, absent soliers = {absent_soldiers}, present soldiers = {present_soldiers}')


    AllNumber = len(helpers.fetchSoldiers())
    AllVacations = helpers.getActiveVacations(with_disabled=True)
    AllVac = len(AllVacations)
    AllAbsent = AllVac
    AllPresent = AllNumber - AllAbsent


    Weekday = days_map[date.today().weekday()]

    
    # print(AllVacations)

    counter = 1
    IterableFieldsDict = []
    for i, tup in enumerate(AllVacations):
        sampleDict = {}
        sampleDict['Soldier_ID'] = tup[0]
        sampleDict['Name'] = tup[1]
        sampleDict['From_Date'] = tup[2]
        if(date.fromisoformat(sampleDict['From_Date']) > date.today()):
            continue
        sampleDict['To_Date'] = tup[3]
        sampleDict['State'] = tup[4]
        sampleDict['Summoned'] = tup[5]
        sampleDict['Level'] = enums.ArmyLevels[int(helpers.getLevelFromID(tup[0]))-1]
        sampleDict["Num"] = counter
        IterableFieldsDict.append(sampleDict)
        counter += 1

    fields_to_replace = {"$DEPT_NAME$": Department_Name,
                        "$DATE_AR$": Date_Arabic,
                        "$WEEKDAY$": Weekday,
                        "$NUM_OFFIC$": number_officers,
                        "$PRES_OFFIC$":present_officers,
                        "$ABS_OFFIC$":absent_officers,
                        "$NUM_CAPS$":number_caps,
                        "$PRES_CAPS$":present_caps,
                        "$ABS_CAPS$":absent_caps,
                        "$NUM_SOLD$":number_soldiers,
                        "$PRES_SOLD$":present_soldiers,
                        "$ABS_SOLD$":absent_soldiers,
                        "$ALL$": AllNumber,
                        "$ALL_P$": AllPresent,
                        "$ALL_A$": AllAbsent,
                        "$ALL_V$": AllVac
                        }
    

    doccc = docx.Document('../pdf/templates/tamam.docx')
    Replace_Placeholders_Inside_Document(doccc, fields_to_replace=fields_to_replace, Iterable_Fields = IterableFieldsDict)
    ConvertAndSave(document=doccc, typeDoc='تمامات', put_watermark=True)
    return True




def WaterMarkPDF(path_to_watermark:str, path_to_pdf:str):
    pdfFile = open(path_to_pdf, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFile)
    # minutesFirstPage = pdfReader.getPage(0)
    pdfWatermarkReader = PyPDF2.PdfReader(open(path_to_watermark, 'rb'))
    # minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
    pdfWriter = PyPDF2.PdfWriter()
    # pdfWriter.addPage(minutesFirstPage)

    for pageNum in range(0, len(pdfReader.pages)):
            pageObj = pdfReader.pages[pageNum]
            pageObj.merge_page(pdfWatermarkReader.pages[0])
            pdfWriter.add_page(pageObj)
    
    pdfFile.close()
    resultPdfFile = open(path_to_pdf, 'wb')
    pdfWriter.write(resultPdfFile)
    resultPdfFile.close()

def add_watermark(input_pdf, watermark_pdf):
    # Read the watermark PDF
    with open(watermark_pdf, "rb") as watermark_file:
        watermark = PyPDF2.PdfReader(watermark_file)
        watermark_page = watermark.pages[0]

        # Read the input PDF
        with open(input_pdf, "rb+") as input_file:
            input_pdf_reader = PyPDF2.PdfReader(input_file)
            output_pdf_writer = PyPDF2.PdfWriter()

            # Iterate through all the pages in the input PDF
            for page_num in range(len(input_pdf_reader.pages)):
                page = input_pdf_reader.pages[page_num]
                # Merge the watermark with the page
                page.merge_page(watermark_page)
                # Add the merged page to the output PDF
                output_pdf_writer.add_page(page)

            # Write the output PDF to a file            
            output_pdf_writer.write(input_pdf)






def ConvertAndSave(document, typeDoc:str, date_of_doc = date.today().isoformat(), put_watermark:bool=False):
    # filePath, extension = os.path.splitext(filePath)
    outputPath =  os.path.join(os.path.split(os.getcwd())[0],typeDoc)
    if(not os.path.isdir(outputPath)):
        os.mkdir(outputPath)

    mapper = {'تمامات': 'تمام_يوم', 'تصاريح':'تصاريح_يوم', 'يوميات_تحركات':'يوميات_تحركات_يوم'}
    
    filepath=os.path.join(outputPath, f'{mapper[typeDoc]}_{date_of_doc}')

    filepath = re.sub(r'\\', '/', filepath)
    

    document.save(f'{filepath}.docx')
    docx2pdf.convert(f'{filepath}.docx', f'{filepath}.pdf', keep_active=True)

    if(put_watermark):
        add_watermark(watermark_pdf='../pdf/templates/watermark.pdf', input_pdf=f'{filepath}.pdf')


    os.startfile(f'{filepath}.pdf')



